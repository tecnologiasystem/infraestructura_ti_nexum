import os
import shutil
import zipfile
import pandas as pd
import logging
import pythoncom
import win32com.client
import pypandoc

from docx import Document
from docx.shared import Inches
from fastapi import UploadFile
from Comun.rutas_archivos import (
    UPLOAD_EXCEL_PREFORMA,
    UPLOAD_CARTAS_IMPULSO,
    UPLOAD_EXCEL_EJECUCION_IMPULSO,
    UPLOAD_FOLDER
)

"""
Configuración del logger para registrar eventos, errores y datos importantes
durante la ejecución de las funciones en este módulo, facilitando el seguimiento
y depuración de procesos.
"""
logger = logging.getLogger("procesal_bll")
logging.basicConfig(level=logging.INFO)

def allowed_file(filename):
    """
    Valida si el archivo tiene una extensión permitida para el procesamiento.
    Retorna True si el archivo tiene extensión '.csv' o '.xlsx', False en otro caso.
    Esto previene procesamiento de archivos no deseados o inseguros.
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {"csv", "xlsx"}

def export_preforma_excel_impulso():
    """
    Obtiene la ruta completa del archivo Excel plantilla para generar cartas de impulso.
    Verifica la existencia del archivo y lanza excepción si no existe.
    Retorna la ruta absoluta para que pueda ser utilizada en descargas o lecturas posteriores.
    """
    file_path = os.path.join(UPLOAD_EXCEL_PREFORMA, "Preforma_ImpulsoProcesal.xlsx")
    if not os.path.exists(file_path):
        raise FileNotFoundError("El archivo Preforma_ImpulsoProcesal.xlsx no existe")
    return file_path

def export_preforma_carta_impulso():
    """
    Similar a la función anterior, pero para el archivo plantilla Word que sirve
    como base para las cartas de impulso.
    """
    file_path = os.path.join(UPLOAD_EXCEL_PREFORMA, "SOLICITUD_AMPLIACION.docx")
    if not os.path.exists(file_path):
        raise FileNotFoundError("El archivo SOLICITUD_AMPLIACION.docx no existe")
    return file_path

def convertir_html_a_docx(html_str, output_file):
    """
    Convierte un string HTML en un archivo Word (.docx).
    - Crea un archivo HTML temporal con la estructura básica y el contenido HTML.
    - Utiliza pypandoc para convertir el HTML temporal a DOCX, usando un documento de referencia para estilos.
    - Maneja errores de conversión y elimina el archivo temporal al finalizar.
    Esto permite transformar contenido HTML personalizado en documentos Word.
    """
    temp_html_path = output_file.replace('.docx', '.html')

    formatted_html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
    <style>@page {{ size: letter; margin: 2.54cm; }}</style>
    </head><body>{html_str}</body></html>"""

    try:
        with open(temp_html_path, 'w', encoding='utf-8') as f:
            f.write(formatted_html)

        reference_doc = os.path.join(UPLOAD_EXCEL_PREFORMA, "SOLICITUD_AMPLIACION.docx")
        pypandoc.convert_file(
            temp_html_path,
            'docx',
            outputfile=output_file,
            extra_args=['--reference-doc=' + reference_doc]
        )
    except Exception as e:
        logger.error(f"Error al convertir HTML a DOCX: {e}")
        raise
    finally:
        if os.path.exists(temp_html_path):
            os.remove(temp_html_path)

def convertir_docx_a_pdf(docx_path, pdf_path):
    """
    Convierte un archivo DOCX en PDF utilizando Microsoft Word mediante automatización COM.
    - Inicializa COM.
    - Abre Word en segundo plano, carga el DOCX y guarda como PDF.
    - Cierra Word y libera recursos COM.
    Esto automatiza la generación de PDFs para facilitar distribución.
    """
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # Código 17 es para formato PDF
        doc.Close()
        word.Quit()
    finally:
        pythoncom.CoUninitialize()

def replace_placeholders_in_docx(docx_file, replacements):
    """
    Reemplaza los marcadores o placeholders en un archivo Word (.docx) con valores específicos.
    - Abre el archivo Word.
    - Recorre todos los párrafos y celdas en tablas buscando las claves a reemplazar.
    - Sustituye las claves por los valores provistos.
    - Guarda el documento modificado.
    Esto permite generar documentos personalizados con datos dinámicos.
    """
    doc = Document(docx_file)
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    doc.save(docx_file)

def guardar_carta_impulso_bll(data, query_args):
    """
    Genera cartas de impulso procesal personalizadas a partir de un contenido HTML base
    y datos obtenidos de un archivo Excel.
    Pasos:
    - Extrae contenido HTML, nombre de archivo Excel y ID de proceso.
    - Verifica existencia y carga Excel con pandas.
    - Filtra filas según tipo de preforma (documento) a procesar.
    - Para cada fila, genera un archivo DOCX personalizado reemplazando placeholders en el HTML.
    - Convierte el DOCX a PDF si se indica, y elimina el DOCX.
    - Acumula las rutas de archivos generados y devuelve un resumen.
    Esto automatiza la creación masiva de documentos legales personalizados.
    """
    html_content = data.get("content", "")
    excel_filename = data.get("excelFilename", "")
    if not excel_filename:
        raise ValueError("No se indicó el nombre del archivo Excel")

    process_id = data.get("processId", "0")

    excel_path = os.path.join(UPLOAD_EXCEL_EJECUCION_IMPULSO, excel_filename)
    print(f"📁 Buscando archivo Excel en: {excel_path}")
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"El archivo Excel {excel_filename} no existe")

    output_type = query_args.get("output", "pdf")
    preforma_index = str(query_args.get("preforma", "1"))
    print(f"🧾 Procesando preforma: preforma_{preforma_index}")

    df = pd.read_excel(excel_path)
    df["TIPO DOCUMENTO"] = df["TIPO DOCUMENTO"].astype(str).str.strip()

    matching = df.loc[df["TIPO DOCUMENTO"].str.contains(f"preforma_{preforma_index}", case=False)]

    if matching.empty:
        print(f"⚠️ No se encontraron filas con TIPO DOCUMENTO=preforma_{preforma_index}. Se procesarán todas.")
        matching = df

    os.makedirs(UPLOAD_CARTAS_IMPULSO, exist_ok=True)
    generated_files = []

    for idx, row in matching.iterrows():
        nuevo_nombre = str(row["Nuevo_Nombre"]).strip()
        final_docx = os.path.join(UPLOAD_CARTAS_IMPULSO, f"{process_id}_{nuevo_nombre}.docx")
        print(f"📝 Generando documento para: {nuevo_nombre}")

        replacements = {
            "{nit}": str(int(row.get("NIT", ""))) if pd.notna(row.get("NIT", "")) else "",
            "{juzgado}": str(row.get("JUZGADO", "")),
            "{nombre}": str(row.get("NOMBRE", "")),
            "{cedula}": str(row.get("CEDULA", "")),
            "{radicado}": str(row.get("RADICADO", "")),
            "{descripcion1}": str(row.get("DESCRIPCION1", "")),
            "{descripcion2}": str(row.get("DESCRIPCION2", ""))
        }

        print(f"🔄 Reemplazos para esta fila: {replacements}")
        print("🧪 HTML antes del reemplazo:")
        print(html_content)

        html_personalizado = html_content
        for clave, valor in replacements.items():
            html_personalizado = html_personalizado.replace(clave, valor)

        print("✅ HTML con reemplazos:")
        print(html_personalizado)

        convertir_html_a_docx(html_personalizado, final_docx)

        if output_type == "pdf":
            pdf_path = final_docx.replace(".docx", ".pdf")
            try:
                convertir_docx_a_pdf(final_docx, pdf_path)
                if os.path.exists(pdf_path):
                    os.remove(final_docx)
                    print(f"✅ PDF generado: {pdf_path}")
                    generated_files.append(pdf_path)
                else:
                    print(f"❌ PDF no encontrado, se deja DOCX: {final_docx}")
                    generated_files.append(final_docx)
            except Exception as e:
                print(f"❌ Error convirtiendo a PDF: {e}")
                generated_files.append(final_docx)
        else:
            print(f"📄 Documento generado: {final_docx}")
            generated_files.append(final_docx)

    return {
        "message": f"Se generaron {len(generated_files)} archivo(s).",
        "files": generated_files,
        "processID": process_id
    }

def exportar_cartas_impulso_bll(process_id):
    """
    Crea un archivo ZIP con todas las cartas generadas de un proceso específico.
    - El nombre del ZIP incorpora el ID del proceso para identificarlo.
    - Recorre la carpeta de cartas y añade al ZIP sólo las que pertenezcan a ese proceso.
    - Devuelve la ruta al archivo ZIP generado para que pueda ser descargado o enviado.
    """
    zip_filename = f"CartasImpulso_{process_id}.zip"
    zip_path = os.path.join(UPLOAD_FOLDER, zip_filename)

    with zipfile.ZipFile(zip_path, "w") as zipf:
        for root, _, files in os.walk(UPLOAD_CARTAS_IMPULSO):
            for f in files:
                if f.startswith(f"{process_id}_"):
                    zipf.write(os.path.join(root, f), f)

    return zip_path

async def upload_excel_juridico_bll(file: UploadFile):
    """
    Función asíncrona para subir archivos Excel jurídicos a la carpeta de ejecución.
    - Valida que el archivo exista y tenga extensión permitida.
    - Crea la carpeta destino si no existe.
    - Guarda el archivo físicamente en disco.
    - Devuelve un mensaje de confirmación y el nombre del archivo subido.
    Esta función se usa en endpoints FastAPI para recepción de archivos desde clientes.
    """
    if not file:
        raise ValueError("No se encontró el archivo.")
    
    if not allowed_file(file.filename):
        raise ValueError("Archivo no permitido.")
    
    filename = file.filename
    upload_folder = UPLOAD_EXCEL_EJECUCION_IMPULSO
    os.makedirs(upload_folder, exist_ok=True)
    filepath = os.path.join(upload_folder, filename)

    with open(filepath, "wb") as buffer:
        buffer.write(await file.read())

    return {"message": f"Archivo Excel {filename} subido correctamente", "filename": filename}
